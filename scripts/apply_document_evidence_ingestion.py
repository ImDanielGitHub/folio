from __future__ import annotations

import ast
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]


def read(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8")


def write(path: str, content: str) -> None:
    target = ROOT / path
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")


def insert_method_before(path: str, class_name: str, before_name: str, method: str) -> None:
    content = read(path)
    tree = ast.parse(content)
    owner = next(node for node in tree.body if isinstance(node, ast.ClassDef) and node.name == class_name)
    before = next(
        node for node in owner.body
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name == before_name
    )
    lines = content.splitlines(keepends=True)
    start = before.lineno - 1
    write(path, "".join(lines[:start]) + method.rstrip() + "\n\n" + "".join(lines[start:]))


MIGRATION = '''    Migration(
        version={version},
        name="document_evidence_ingestion",
        sql="""
        CREATE TABLE document_sources (
            document_id TEXT PRIMARY KEY,
            workspace_id TEXT NOT NULL REFERENCES workspaces(workspace_id),
            filename TEXT NOT NULL CHECK (length(trim(filename)) BETWEEN 1 AND 255),
            media_type TEXT NOT NULL,
            source_sha256 TEXT NOT NULL CHECK (length(source_sha256) = 64),
            size_bytes INTEGER NOT NULL CHECK (size_bytes > 0),
            page_count INTEGER,
            character_count INTEGER NOT NULL CHECK (character_count >= 0),
            extraction_version TEXT NOT NULL,
            source_bytes BLOB NOT NULL,
            received_at TEXT NOT NULL,
            UNIQUE (workspace_id, source_sha256)
        );

        CREATE INDEX document_sources_workspace_time
            ON document_sources(workspace_id, received_at, document_id);
        """,
    ),
'''

DOCUMENTS = '''"""Bounded local document extraction into immutable evidence and knowledge records."""

from __future__ import annotations

import hashlib
import html
import json
import re
import zipfile
from dataclasses import dataclass
from datetime import UTC, datetime
from html.parser import HTMLParser
from io import BytesIO
from pathlib import PurePath
from typing import Any, Literal

from docx import Document
from pypdf import PdfReader

from finance_agent.storage import SQLiteStore, canonical_json

MAX_DOCUMENT_BYTES = 20_000_000
MAX_PDF_PAGES = 250
MAX_EXTRACTED_CHARACTERS = 500_000
EXTRACTION_VERSION = "document_extract@1"
DocumentKind = Literal[
    "receipt", "invoice", "contract", "bank_statement", "tax_document",
    "correspondence", "note", "other",
]


class DocumentIngestError(ValueError):
    pass


def _stable_id(prefix: str, *parts: str) -> str:
    digest = hashlib.sha256("\\0".join(parts).encode()).hexdigest()[:24]
    return f"{prefix}_{digest}"


def _clean_filename(filename: str) -> str:
    value = PurePath(filename.replace("\\", "/")).name.strip()
    value = re.sub(r"[\\x00-\\x1f\\x7f]", "", value)
    if not value:
        raise DocumentIngestError("document filename is empty")
    return value[:255]


def _bound_text(value: str) -> str:
    value = value.replace("\\x00", " ")
    value = re.sub(r"[ \\t]+", " ", value)
    value = re.sub(r"\\n{3,}", "\\n\\n", value).strip()
    if not value:
        raise DocumentIngestError("document contains no extractable text")
    if len(value) > MAX_EXTRACTED_CHARACTERS:
        raise DocumentIngestError(
            f"document text exceeds the {MAX_EXTRACTED_CHARACTERS} character limit"
        )
    return value


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.suppressed = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in {"script", "style", "template", "noscript"}:
            self.suppressed += 1
        elif tag.lower() in {"p", "br", "li", "tr", "h1", "h2", "h3", "h4"}:
            self.parts.append("\\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "template", "noscript"} and self.suppressed:
            self.suppressed -= 1
        elif tag.lower() in {"p", "li", "tr"}:
            self.parts.append("\\n")

    def handle_data(self, data: str) -> None:
        if not self.suppressed:
            self.parts.append(data)


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    filename: str
    media_type: str
    text: str
    page_count: int | None
    metadata: dict[str, Any]


def extract_document(filename: str, content: bytes) -> ExtractedDocument:
    clean_name = _clean_filename(filename)
    if not content:
        raise DocumentIngestError("document is empty")
    if len(content) > MAX_DOCUMENT_BYTES:
        raise DocumentIngestError(
            f"document exceeds the {MAX_DOCUMENT_BYTES} byte limit"
        )
    suffix = PurePath(clean_name).suffix.lower()
    if content.startswith(b"%PDF-"):
        try:
            reader = PdfReader(BytesIO(content), strict=True)
        except Exception as exc:
            raise DocumentIngestError("PDF could not be parsed") from exc
        if reader.is_encrypted:
            raise DocumentIngestError("encrypted PDFs are not accepted")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise DocumentIngestError(
                f"PDF exceeds the {MAX_PDF_PAGES} page limit"
            )
        parts: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            try:
                parts.append(page.extract_text() or "")
            except Exception as exc:
                raise DocumentIngestError(
                    f"PDF page {index} could not be extracted"
                ) from exc
        return ExtractedDocument(
            clean_name,
            "application/pdf",
            _bound_text("\\n\\n".join(parts)),
            len(reader.pages),
            {"ocrUsed": False, "encrypted": False},
        )
    if content.startswith(b"PK\\x03\\x04") and suffix == ".docx":
        try:
            with zipfile.ZipFile(BytesIO(content)) as archive:
                names = set(archive.namelist())
                if any(name.lower().endswith("vbaproject.bin") for name in names):
                    raise DocumentIngestError("macro-enabled Office documents are not accepted")
                if "word/document.xml" not in names:
                    raise DocumentIngestError("DOCX archive has no document body")
        except zipfile.BadZipFile as exc:
            raise DocumentIngestError("DOCX archive is invalid") from exc
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise DocumentIngestError("DOCX could not be parsed") from exc
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return ExtractedDocument(
            clean_name,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _bound_text("\\n".join(parts)),
            None,
            {"ocrUsed": False, "macrosPresent": False},
        )
    if suffix in {".html", ".htm"}:
        try:
            source = content.decode("utf-8", errors="strict")
        except UnicodeDecodeError as exc:
            raise DocumentIngestError("HTML document must be UTF-8") from exc
        parser = _TextHTMLParser()
        parser.feed(source)
        return ExtractedDocument(
            clean_name,
            "text/html",
            _bound_text(html.unescape("".join(parser.parts))),
            None,
            {"ocrUsed": False, "activeContentExecuted": False},
        )
    if suffix in {".txt", ".md"}:
        try:
            source = content.decode("utf-8", errors="strict")
        except UnicodeDecodeError as exc:
            raise DocumentIngestError("text document must be UTF-8") from exc
        return ExtractedDocument(
            clean_name,
            "text/markdown" if suffix == ".md" else "text/plain",
            _bound_text(source),
            None,
            {"ocrUsed": False},
        )
    raise DocumentIngestError(
        "supported document types are PDF, DOCX, HTML, Markdown and UTF-8 text"
    )


class DocumentEvidenceService:
    def __init__(self, store: SQLiteStore) -> None:
        self.store = store

    def ingest(
        self,
        *,
        workspace_id: str,
        filename: str,
        content: bytes,
        document_kind: DocumentKind,
        task_scope: str,
    ) -> dict[str, object]:
        if document_kind not in {
            "receipt", "invoice", "contract", "bank_statement", "tax_document",
            "correspondence", "note", "other",
        }:
            raise DocumentIngestError("unsupported document kind")
        scope = task_scope.strip()
        if not scope:
            raise DocumentIngestError("task scope must not be blank")
        digest = hashlib.sha256(content).hexdigest()
        existing = self.store.fetch_one(
            """
            SELECT document_id, filename, media_type, size_bytes, page_count,
                   character_count, received_at
            FROM document_sources
            WHERE workspace_id = ? AND source_sha256 = ?
            """,
            (workspace_id, digest),
        )
        if existing is not None:
            return {
                "documentId": str(existing["document_id"]),
                "status": "deduplicated",
                "filename": str(existing["filename"]),
                "mediaType": str(existing["media_type"]),
                "sizeBytes": int(existing["size_bytes"]),
                "pageCount": existing["page_count"],
                "characterCount": int(existing["character_count"]),
                "receivedAt": str(existing["received_at"]),
                "sourceSha256": digest,
            }
        extracted = extract_document(filename, content)
        received_at = datetime.now(UTC).isoformat()
        document_id = _stable_id("doc", workspace_id, digest)
        evidence_id = _stable_id("evd", workspace_id, digest)
        metadata = {
            **extracted.metadata,
            "filename": extracted.filename,
            "mediaType": extracted.media_type,
            "sourceSha256": digest,
            "sizeBytes": len(content),
            "pageCount": extracted.page_count,
            "extractionVersion": EXTRACTION_VERSION,
        }
        record_value = {
            "documentId": document_id,
            "workspaceId": workspace_id,
            "documentKind": document_kind,
            "title": extracted.filename,
            "taskScope": scope,
            "sourceKind": "document",
            "sourceRef": document_id,
            "evidenceId": evidence_id,
            "receivedAt": received_at,
            "extractedText": extracted.text,
            "contentHash": digest,
            "metadata": metadata,
        }
        record_hash = hashlib.sha256(canonical_json(record_value).encode()).hexdigest()
        with self.store.transaction() as connection:
            connection.execute(
                """
                INSERT INTO document_sources(
                    document_id, workspace_id, filename, media_type, source_sha256,
                    size_bytes, page_count, character_count, extraction_version,
                    source_bytes, received_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    document_id,
                    workspace_id,
                    extracted.filename,
                    extracted.media_type,
                    digest,
                    len(content),
                    extracted.page_count,
                    len(extracted.text),
                    EXTRACTION_VERSION,
                    content,
                    received_at,
                ),
            )
            connection.execute(
                """
                INSERT INTO evidence_links(
                    evidence_id, workspace_id, source_item_id, source_row_id,
                    label, created_at
                ) VALUES (?, ?, NULL, NULL, ?, ?)
                """,
                (
                    evidence_id,
                    workspace_id,
                    f"{extracted.filename} ({document_kind})",
                    received_at,
                ),
            )
            connection.execute(
                """
                INSERT INTO knowledge_documents(
                    document_id, workspace_id, document_kind, title, task_scope,
                    source_kind, source_ref, source_turn_id, evidence_id, received_at,
                    effective_from, effective_until, extracted_text, content_hash,
                    metadata_json, record_hash
                ) VALUES (?, ?, ?, ?, ?, 'document', ?, NULL, ?, ?, NULL, NULL, ?, ?, ?, ?)
                """,
                (
                    document_id,
                    workspace_id,
                    document_kind,
                    extracted.filename,
                    scope,
                    document_id,
                    evidence_id,
                    received_at,
                    extracted.text,
                    digest,
                    canonical_json(metadata),
                    record_hash,
                ),
            )
        return {
            "documentId": document_id,
            "evidenceId": evidence_id,
            "status": "ingested",
            "filename": extracted.filename,
            "mediaType": extracted.media_type,
            "sizeBytes": len(content),
            "pageCount": extracted.page_count,
            "characterCount": len(extracted.text),
            "receivedAt": received_at,
            "sourceSha256": digest,
            "ocrUsed": False,
        }

    def list_documents(self, workspace_id: str) -> list[dict[str, object]]:
        rows = self.store.fetch_all(
            """
            SELECT document_id, filename, media_type, source_sha256, size_bytes,
                   page_count, character_count, extraction_version, received_at
            FROM document_sources WHERE workspace_id = ?
            ORDER BY received_at DESC, document_id DESC
            """,
            (workspace_id,),
        )
        return [
            {
                "documentId": str(row["document_id"]),
                "filename": str(row["filename"]),
                "mediaType": str(row["media_type"]),
                "sourceSha256": str(row["source_sha256"]),
                "sizeBytes": int(row["size_bytes"]),
                "pageCount": row["page_count"],
                "characterCount": int(row["character_count"]),
                "extractionVersion": str(row["extraction_version"]),
                "receivedAt": str(row["received_at"]),
            }
            for row in rows
        ]
'''

SERVICE_METHODS = '''    async def ingest_document(
        self,
        *,
        workspace_id: str,
        filename: str,
        content: bytes,
        document_kind: str,
        task_scope: str,
    ) -> Mapping[str, object]:
        if workspace_id != WORKSPACE_ID or self._workspace_destroyed:
            raise KeyError(workspace_id)
        async with self._lock:
            result = DocumentEvidenceService(self.store).ingest(
                workspace_id=workspace_id,
                filename=filename,
                content=content,
                document_kind=document_kind,  # type: ignore[arg-type]
                task_scope=task_scope,
            )
            self.working_understanding.ensure_current(workspace_id=workspace_id)
        return result

    async def list_documents(
        self, *, workspace_id: str
    ) -> tuple[Mapping[str, object], ...]:
        if workspace_id != WORKSPACE_ID or self._workspace_destroyed:
            raise KeyError(workspace_id)
        return tuple(DocumentEvidenceService(self.store).list_documents(workspace_id))
'''

ROUTES = '''    @router.post("/v1/workspaces/{workspace_id}/documents", status_code=201)
    async def ingest_document(
        workspace_id: PathIdentifier,
        services: Services,
        document_kind: Annotated[
            str,
            Form(
                alias="documentKind",
                pattern=r"^(receipt|invoice|contract|bank_statement|tax_document|correspondence|note|other)$",
            ),
        ],
        task_scope: Annotated[
            str,
            Form(alias="taskScope", min_length=1, max_length=200),
        ],
        file: Annotated[UploadFile, File()],
    ) -> dict[str, object]:
        try:
            content = await read_upload_with_limit(file, max_bytes=20_000_000)
            return dict(
                await services.ingest_document(
                    workspace_id=workspace_id,
                    filename=file.filename or "document",
                    content=content,
                    document_kind=document_kind,
                    task_scope=task_scope,
                )
            )
        except UploadTooLarge as exc:
            raise HTTPException(status_code=413, detail=str(exc)) from exc
        except DocumentIngestError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc

    @router.get("/v1/workspaces/{workspace_id}/documents")
    async def list_documents(
        workspace_id: PathIdentifier,
        services: Services,
    ) -> dict[str, object]:
        documents = await services.list_documents(workspace_id=workspace_id)
        return {"workspaceId": workspace_id, "documents": list(documents)}

'''

TESTS = '''from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen.canvas import Canvas

from finance_agent.api.services import LocalRouteServices
from finance_agent.storage.documents import DocumentEvidenceService, DocumentIngestError


def pdf_bytes(text: str) -> bytes:
    buffer = BytesIO()
    canvas = Canvas(buffer)
    canvas.drawString(72, 720, text)
    canvas.save()
    return buffer.getvalue()


def docx_bytes(text: str) -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.asyncio
async def test_pdf_ingestion_is_hashed_indexed_and_idempotent(tmp_path: Path) -> None:
    services = LocalRouteServices(tmp_path / "folio.sqlite3", auto_seed=True)
    content = pdf_bytes("Client invoice INV-123 for NZD 115.00")
    first = await services.ingest_document(
        workspace_id="ws_koru_studio",
        filename="invoice.pdf",
        content=content,
        document_kind="invoice",
        task_scope="gst preparation",
    )
    second = await services.ingest_document(
        workspace_id="ws_koru_studio",
        filename="invoice.pdf",
        content=content,
        document_kind="invoice",
        task_scope="gst preparation",
    )
    assert first["status"] == "ingested"
    assert second["status"] == "deduplicated"
    assert first["documentId"] == second["documentId"]
    assert first["ocrUsed"] is False
    fts = services.store.fetch_one(
        "SELECT body FROM knowledge_fts WHERE record_type = 'document' AND record_id = ?",
        (first["documentId"],),
    )
    assert fts is not None
    assert "INV-123" in str(fts["body"])
    source = services.store.fetch_one(
        "SELECT source_bytes FROM document_sources WHERE document_id = ?",
        (first["documentId"],),
    )
    assert bytes(source["source_bytes"]) == content
    await services.aclose()


@pytest.mark.asyncio
async def test_docx_and_utf8_text_are_supported(tmp_path: Path) -> None:
    services = LocalRouteServices(tmp_path / "folio.sqlite3", auto_seed=True)
    docx = await services.ingest_document(
        workspace_id="ws_koru_studio",
        filename="contract.docx",
        content=docx_bytes("Signed client services contract"),
        document_kind="contract",
        task_scope="customer context",
    )
    text = await services.ingest_document(
        workspace_id="ws_koru_studio",
        filename="note.txt",
        content="Owner note about the Acme invoice".encode(),
        document_kind="note",
        task_scope="cash forecast",
    )
    assert docx["mediaType"].endswith("document")
    assert text["mediaType"] == "text/plain"
    assert len(await services.list_documents(workspace_id="ws_koru_studio")) == 2
    await services.aclose()


def test_encrypted_pdf_and_macro_enabled_office_file_fail_closed(tmp_path: Path) -> None:
    from finance_agent.finance import FinanceEngine
    from finance_agent.storage import SQLiteStore

    store = SQLiteStore(tmp_path / "folio.sqlite3")
    FinanceEngine(store).initialise()
    evidence = DocumentEvidenceService(store)
    source = PdfReader(BytesIO(pdf_bytes("secret")))
    writer = PdfWriter()
    for page in source.pages:
        writer.add_page(page)
    writer.encrypt("password")
    encrypted = BytesIO()
    writer.write(encrypted)
    with pytest.raises(DocumentIngestError, match="encrypted PDFs"):
        evidence.ingest(
            workspace_id="ws_koru_studio",
            filename="secret.pdf",
            content=encrypted.getvalue(),
            document_kind="other",
            task_scope="test",
        )

    macro = BytesIO()
    with ZipFile(macro, "w", ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", "<document />")
        archive.writestr("word/vbaProject.bin", b"macro")
    with pytest.raises(DocumentIngestError, match="macro-enabled"):
        from finance_agent.storage.documents import extract_document

        extract_document("unsafe.docx", macro.getvalue())
'''


def add_dependencies() -> None:
    path = "services/api/pyproject.toml"
    content = read(path)
    marker = '  "jsonschema>=4.25,<5",\n'
    additions = '  "pypdf>=6,<7",\n  "python-docx>=1.2,<2",\n'
    if '"pypdf>=' not in content:
        if marker not in content:
            raise RuntimeError("dependency marker missing")
        content = content.replace(marker, marker + additions, 1)
        write(path, content)


def add_migration() -> None:
    path = "services/api/src/finance_agent/storage/migrations.py"
    content = read(path)
    versions = [int(value) for value in re.findall(r"version=(\d+)", content)]
    version = max(versions) + 1
    closing = content.rfind("\n)")
    if closing < 0:
        raise RuntimeError("MIGRATIONS tuple close not found")
    prefix = content[:closing].rstrip()
    if not prefix.endswith(","):
        prefix += ","
    write(path, prefix + "\n" + MIGRATION.format(version=version) + content[closing:])


def add_module_and_service() -> None:
    write("services/api/src/finance_agent/storage/documents.py", DOCUMENTS)
    path = "services/api/src/finance_agent/api/services.py"
    content = read(path)
    marker = "from finance_agent.storage.encrypted_exports import decrypt_backup, encrypt_backup\n"
    import_line = "from finance_agent.storage.documents import DocumentEvidenceService\n"
    if import_line not in content:
        if marker not in content:
            raise RuntimeError("encrypted export import marker missing")
        content = content.replace(marker, marker + import_line, 1)
        write(path, content)
    insert_method_before(path, "LocalRouteServices", "set_gst_mapping", SERVICE_METHODS)


def update_protocol_and_routes() -> None:
    path = "services/api/src/finance_agent/api/routes/dependencies.py"
    content = read(path)
    marker = "    async def set_gst_mapping(\n"
    addition = '''    async def ingest_document(\n        self, *, workspace_id: str, filename: str, content: bytes,\n        document_kind: str, task_scope: str\n    ) -> Mapping[str, object]: ...\n\n    async def list_documents(\n        self, *, workspace_id: str\n    ) -> tuple[Mapping[str, object], ...]: ...\n\n'''
    if marker not in content:
        raise RuntimeError("GST protocol marker missing")
    content = content.replace(marker, addition + marker, 1)
    write(path, content)

    path = "services/api/src/finance_agent/api/routes/router.py"
    content = read(path)
    import_marker = "from finance_agent.connectors.base import ConnectorError\n"
    import_line = "from finance_agent.storage.documents import DocumentIngestError\n"
    if import_line not in content:
        if import_marker not in content:
            raise RuntimeError("connector import marker missing")
        content = content.replace(import_marker, import_marker + import_line, 1)
    route_marker = '    @router.post("/v1/workspaces/{workspace_id}/accounting/gst-mappings")\n'
    if route_marker not in content:
        raise RuntimeError("GST route marker missing")
    content = content.replace(route_marker, ROUTES + route_marker, 1)
    write(path, content)


def add_tests_and_docs() -> None:
    write("services/api/tests/storage/test_document_evidence_ingestion.py", TESTS)
    write("docs/DOCUMENT_INGESTION.md", '''# Document ingestion boundary\n\nFolio accepts bounded local PDF, DOCX, HTML, Markdown and UTF-8 text files. It retains the source bytes, SHA-256 digest, extraction metadata, evidence identifier and extracted text in the local workspace. Duplicate source bytes are idempotent. Extracted text enters the existing source-linked knowledge index and is never treated as finance truth merely because it appeared in a document.\n\nEncrypted PDFs, macro-enabled Office documents, unsupported formats, over-sized documents and documents without extractable text fail closed. HTML active content is never executed. PDF extraction is text-layer only; no OCR is performed or claimed. Owners must inspect source evidence before accepting a classification, tax treatment or business fact derived from a document.\n''')
    path = "docs/AUDIT_PROGRAMME.md"
    content = read(path)
    section = '''\n\n## Stack 12: bounded document evidence ingestion\n\n- Local PDF, DOCX, HTML, Markdown and UTF-8 text sources are accepted within byte, page and character limits.\n- Source bytes, SHA-256, extraction version, evidence ID and extracted text are retained locally.\n- Duplicate bytes are idempotent and extracted text enters the provenance-backed knowledge index.\n- Encrypted PDFs, Office macros, unsupported formats and empty extraction fail closed.\n- HTML active content is not executed.\n- OCR is not performed or claimed.\n'''
    if "## Stack 12: bounded document evidence ingestion" not in content:
        write(path, content.rstrip() + section + "\n")


def main() -> None:
    add_dependencies()
    add_migration()
    add_module_and_service()
    update_protocol_and_routes()
    add_tests_and_docs()
    print("document evidence ingestion changes applied")


if __name__ == "__main__":
    main()
